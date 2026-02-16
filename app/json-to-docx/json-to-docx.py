"""
Stage 5: JSON → DOCX 変換
final_report_*.json を読み込み、Word文書を生成する。

Usage:
    uv run app/json-to-docx/json-to-docx.py \
        -i data/final-output/final-report-per-company/final_report_12044.json \
        -o data/final-output/docx-per-company/report_12044.docx
"""

import argparse
import json
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# スタイル設定
# ---------------------------------------------------------------------------

FONT_NAME_JA = "游ゴシック"
FONT_NAME_EN = "Calibri"
COLOR_DARK = RGBColor(0x1A, 0x1A, 0x2E)
COLOR_ACCENT = RGBColor(0x00, 0x56, 0x8B)
COLOR_LIGHT_BG = RGBColor(0xF2, 0xF5, 0xF8)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color: RGBColor):
    """セルの背景色を設定する。"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): f"{color}"},
    )
    shading.append(shading_elem)


def fmt_run(run, size=10.5, bold=False, color=None):
    """Runのフォント設定をまとめて行う。"""
    run.font.size = Pt(size)
    run.font.name = FONT_NAME_EN
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_JA)
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text, style=None, size=10.5, bold=False, color=None,
                  alignment=None, space_after=Pt(6)):
    """段落を追加するヘルパー。"""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    fmt_run(run, size=size, bold=bold, color=color)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """テーブルを追加するヘルパー。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        fmt_run(run, size=9, bold=True, color=COLOR_WHITE)
        set_cell_shading(cell, COLOR_ACCENT)

    # データ行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            fmt_run(run, size=9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, COLOR_LIGHT_BG)

    # カラム幅
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # テーブル後のスペース
    return table


def fmt_number(val):
    """数値を読みやすい文字列にフォーマットする。"""
    if val is None:
        return "—"
    if isinstance(val, float):
        if abs(val) >= 1_000_000:
            return f"{val:,.0f}"
        return f"{val:,.2f}"
    if isinstance(val, int):
        return f"{val:,}"
    return str(val)


# ---------------------------------------------------------------------------
# セクション描画
# ---------------------------------------------------------------------------

def render_executive_summary(doc, section):
    doc.add_heading(section["title"], level=1)
    add_paragraph(doc, section.get("content", ""), size=10.5)


def render_company_overview(doc, section):
    doc.add_heading(section["title"], level=1)

    for sub in section.get("subsections", []):
        doc.add_heading(sub["title"], level=2)
        content = sub.get("content", {})

        if sub["id"] == "company_and_region":
            render_company_and_region(doc, content)
        elif sub["id"] == "industry_environment":
            render_industry_environment(doc, content)


def render_company_and_region(doc, content):
    # 企業情報テーブル
    info = content.get("企業情報", {})
    if info:
        add_paragraph(doc, "企業情報", size=10.5, bold=True, color=COLOR_ACCENT)
        rows = [[k, fmt_number(v)] for k, v in info.items()]
        add_table(doc, ["項目", "内容"], rows, col_widths=[5, 10])

    # 地域特性テキスト
    feature_keys = [
        "事業・営業・受注戦略の地域的特徴",
        "人的資本の地域的特徴",
        "財務構造の地域的特徴",
        "全体の地域的特徴",
    ]
    for key in feature_keys:
        text = content.get(key, "")
        if text:
            add_paragraph(doc, key, size=10.5, bold=True, color=COLOR_ACCENT)
            add_paragraph(doc, text, size=10)


def render_industry_environment(doc, content):
    # PEST分析
    pest = content.get("pest", {})
    if pest:
        add_paragraph(doc, "PEST分析", size=10.5, bold=True, color=COLOR_ACCENT)
        label_map = {
            "political": "Political（政治）",
            "economic": "Economic（経済）",
            "social": "Social（社会）",
            "technological": "Technological（技術）",
        }
        rows = [[label_map.get(k, k), v] for k, v in pest.items()]
        add_table(doc, ["要因", "概要"], rows, col_widths=[4, 12])

    # 需給動向
    ds = content.get("demand_supply", {})
    if ds:
        add_paragraph(doc, "需給動向", size=10.5, bold=True, color=COLOR_ACCENT)
        rows = [[k, v] for k, v in ds.items()]
        add_table(doc, ["区分", "概要"], rows, col_widths=[3, 13])

    # 資本市場
    cm = content.get("capital_market", {})
    if cm:
        add_paragraph(doc, "資本市場", size=10.5, bold=True, color=COLOR_ACCENT)
        rows = [[k, v] for k, v in cm.items()]
        add_table(doc, ["区分", "概要"], rows, col_widths=[4, 12])

    # シナリオ
    sc = content.get("scenarios", {})
    if sc:
        add_paragraph(doc, "シナリオ分析", size=10.5, bold=True, color=COLOR_ACCENT)
        rows = [[k, v] for k, v in sc.items()]
        add_table(doc, ["シナリオ", "概要"], rows, col_widths=[3, 13])


def render_analysis(doc, section):
    doc.add_heading(section["title"], level=1)
    content = section.get("content", {})

    # タグ別平均スコア
    tag_avg = content.get("tag_averages", {})
    if tag_avg:
        add_paragraph(doc, "タグ別平均スコア", size=10.5, bold=True, color=COLOR_ACCENT)
        sorted_tags = sorted(tag_avg.items(), key=lambda x: x[1])
        rows = [[tag, f"{score:.2f}"] for tag, score in sorted_tags]
        add_table(doc, ["タグ", "平均スコア"], rows, col_widths=[8, 4])

    # 財務分析
    fin = content.get("financial_analysis", {})
    if fin:
        render_financial_analysis(doc, fin)

    # 定性タグ
    for tag_data in content.get("qualitative_tags", []):
        render_qualitative_tag(doc, tag_data)

    # 強み・弱みの総括
    overall = content.get("overall_summary", {})
    if overall:
        doc.add_heading("総括：強みと課題", level=2)
        strengths = overall.get("strengths", "")
        if strengths:
            add_paragraph(doc, "■ 強み", size=10.5, bold=True, color=COLOR_ACCENT)
            add_paragraph(doc, strengths, size=10)
        weaknesses = overall.get("weaknesses", "")
        if weaknesses:
            add_paragraph(doc, "■ 弱み・課題", size=10.5, bold=True, color=COLOR_ACCENT)
            add_paragraph(doc, weaknesses, size=10)


def render_financial_analysis(doc, fin):
    doc.add_heading("財務分析", level=2)

    # スコア概要
    add_paragraph(
        doc,
        f"タグ: {fin.get('tag', '')}　|　平均スコア: {fin.get('avg_score', 0):.2f}",
        size=10, bold=True,
    )

    # 項目スコア
    items = fin.get("items", {})
    if items:
        add_paragraph(doc, "項目別スコア", size=10, bold=True, color=COLOR_ACCENT)
        rows = [[k, str(v)] for k, v in items.items()]
        add_table(doc, ["項目", "スコア"], rows, col_widths=[8, 4])

    # サマリー
    summary = fin.get("summary", "")
    if summary:
        add_paragraph(doc, summary, size=10)

    # 年度別指標テーブル
    metrics = fin.get("key_metrics", [])
    if metrics:
        add_paragraph(doc, "年度別財務指標", size=10.5, bold=True, color=COLOR_ACCENT)
        for year_data in metrics:
            year = year_data.get("YEAR", "")
            add_paragraph(doc, f"{year}年度", size=10, bold=True)

            categories = [
                "収益性指標", "成長性指標", "コスト構造・固定費分析",
                "効率性指標", "安全性・財務健全性",
                "キャッシュフロー関連指標", "建設業特有指標",
            ]
            rows = []
            for cat in categories:
                cat_data = year_data.get(cat)
                if not cat_data:
                    continue
                for k, v in cat_data.items():
                    rows.append([cat, k, fmt_number(v)])

            if rows:
                add_table(doc, ["カテゴリ", "指標", "値"], rows, col_widths=[5, 6, 4])


def render_qualitative_tag(doc, tag_data):
    tag_name = tag_data.get("tag", "")
    avg = tag_data.get("avg_score", 0)
    doc.add_heading(f"{tag_name}（平均: {avg:.2f}）", level=2)

    # 項目スコア
    items = tag_data.get("items", {})
    if items:
        rows = [[k, str(v)] for k, v in items.items()]
        add_table(doc, ["項目", "スコア"], rows, col_widths=[8, 4])

    # サマリー
    summary = tag_data.get("summary", "")
    if summary:
        add_paragraph(doc, summary, size=10)


def render_strategy(doc, section):
    doc.add_heading(section["title"], level=1)

    for sub in section.get("subsections", []):
        doc.add_heading(sub["title"], level=2)
        content = sub.get("content", {})

        if sub["id"] == "initiatives":
            render_initiatives(doc, content)
        elif sub["id"] == "fit_to_company":
            render_fit_to_company(doc, content)


def render_initiatives(doc, content):
    for init in content.get("initiatives", []):
        name = init.get("施策名", "")
        priority = init.get("priority", "")
        score = init.get("relevance_score", "")

        doc.add_heading(f"施策{priority}: {name}", level=3)
        add_paragraph(doc, f"適合度スコア: {score}/5", size=10, bold=True)

        fields = [
            ("課題適合理由", "課題適合理由"),
            ("地域性適合理由", "地域性適合理由"),
            ("業界特性適合理由", "業界特性適合理由"),
            ("期待効果", "expected_impact"),
        ]
        for label, key in fields:
            val = init.get(key, "")
            if val:
                add_paragraph(doc, f"■ {label}", size=10, bold=True, color=COLOR_ACCENT)
                add_paragraph(doc, val, size=10)

        weaknesses = init.get("対応する弱点", [])
        if weaknesses:
            add_paragraph(doc, f"■ 対応する弱点", size=10, bold=True, color=COLOR_ACCENT)
            for w in weaknesses:
                add_paragraph(doc, f"・{w}", size=10)


def render_fit_to_company(doc, content):
    # 適合理由テキスト
    narrative = content.get("fit_narrative", "")
    if narrative:
        add_paragraph(doc, "適合性の根拠", size=10.5, bold=True, color=COLOR_ACCENT)
        for para_text in narrative.split("\n"):
            para_text = para_text.strip()
            if para_text:
                add_paragraph(doc, para_text, size=10)

    # 他社比較
    comparisons = content.get("comparison", [])
    if comparisons:
        add_paragraph(doc, "他社事例との比較", size=10.5, bold=True, color=COLOR_ACCENT)
        rows = [
            [c.get("企業名", ""), c.get("証券コード", ""), c.get("施策として読み取れる具体的記述", "")]
            for c in comparisons
        ]
        add_table(doc, ["企業名", "証券コード", "具体的記述"], rows, col_widths=[3, 2.5, 10.5])

    # 強み
    strengths = content.get("company_specific_strengths", [])
    if strengths:
        add_paragraph(doc, "当社固有の強み", size=10.5, bold=True, color=COLOR_ACCENT)
        for s in strengths:
            add_paragraph(doc, f"・{s}", size=10)

    # 制約
    constraints = content.get("company_specific_constraints", [])
    if constraints:
        add_paragraph(doc, "当社固有の制約", size=10.5, bold=True, color=COLOR_ACCENT)
        for c in constraints:
            add_paragraph(doc, f"・{c}", size=10)


def render_impact(doc, section):
    doc.add_heading(section["title"], level=1)
    content = section.get("content", {})

    # 前提条件
    assumptions = content.get("assumptions", [])
    if assumptions:
        add_paragraph(doc, "前提条件", size=10.5, bold=True, color=COLOR_ACCENT)
        for a in assumptions:
            add_paragraph(doc, f"・{a}", size=10)

    note = content.get("conservativeness_note", "")
    if note:
        add_paragraph(doc, note, size=9.5, bold=False, color=RGBColor(0x66, 0x66, 0x66))

    # 定量的効果
    qi = content.get("quantitative_impact", {})
    if qi:
        add_paragraph(doc, "定量的効果", size=10.5, bold=True, color=COLOR_ACCENT)
        rows = [
            ["売上高", qi.get("revenue", "")],
            ["営業利益率", qi.get("profit_margin", "")],
            ["キャッシュフロー", qi.get("cash_flow", "")],
        ]
        add_table(doc, ["指標", "見通し"], rows, col_widths=[4, 12])

        calc = qi.get("calculation_notes", "")
        if calc:
            add_paragraph(doc, f"算出根拠: {calc}", size=9.5, color=RGBColor(0x66, 0x66, 0x66))

    # 定性的効果
    ql = content.get("qualitative_impact", {})
    if ql:
        add_paragraph(doc, "定性的効果", size=10.5, bold=True, color=COLOR_ACCENT)
        for e in ql.get("effects", []):
            add_paragraph(doc, f"・{e}", size=10)
        narrative = ql.get("narrative", "")
        if narrative:
            add_paragraph(doc, narrative, size=10)


def render_roadmap(doc, section):
    doc.add_heading(section["title"], level=1)
    content = section.get("content", {})

    phases = [
        ("short_term", "短期（0–1年）"),
        ("mid_term", "中期（1–3年）"),
        ("long_term", "長期（3–5年）"),
    ]

    rows = []
    for key, label in phases:
        phase = content.get(key, {})
        actions = phase.get("actions", [])
        ideal = phase.get("ideal_state", "")
        rows.append([label, "\n".join(f"・{a}" for a in actions), ideal])

    if rows:
        add_table(doc, ["フェーズ", "アクション", "到達目標"], rows, col_widths=[3, 8, 5])


def render_risks(doc, section):
    doc.add_heading(section["title"], level=1)
    content = section.get("content", {})

    risks = content.get("risks", [])
    if not risks:
        return

    rows = []
    for r in risks:
        rows.append([
            r.get("risk_type", ""),
            r.get("risk", ""),
            r.get("mitigation", ""),
            r.get("trigger_or_signal", ""),
        ])

    add_table(
        doc,
        ["リスク種別", "リスク内容", "対応策", "トリガー/シグナル"],
        rows,
        col_widths=[2.5, 5, 5, 4],
    )


# ---------------------------------------------------------------------------
# ドキュメント組み立て
# ---------------------------------------------------------------------------

CLOSING_MESSAGE = (
    "本レポートは、貴社の有価証券報告書をはじめとする公開情報をもとに、"
    "AI分析エンジンによる定量・定性の両面から経営課題の抽出と成長施策の立案を行ったものでございます。\n\n"
    "ここに示した分析結果および施策案が、貴社の経営戦略のご検討における一助となれば幸いです。"
    "なお、本資料はあくまで分析に基づく提案であり、最終的なご判断は貴社の経営環境・方針を踏まえてご決定いただければと存じます。\n\n"
    "ご不明点やさらなる深掘りのご要望がございましたら、お気軽にお申し付けください。"
    "貴社のさらなるご発展を心よりお祈り申し上げます。"
)


def render_closing(doc):
    """締めの挨拶を描画する。"""
    doc.add_page_break()
    doc.add_heading("おわりに", level=1)
    for para_text in CLOSING_MESSAGE.split("\n\n"):
        add_paragraph(doc, para_text.strip(), size=10.5)


SECTION_RENDERERS = {
    "executive_summary": render_executive_summary,
    "company_overview": render_company_overview,
    "analysis": render_analysis,
    "strategy": render_strategy,
    "impact": render_impact,
    "roadmap": render_roadmap,
    "risks": render_risks,
}


def build_document(data: dict) -> Document:
    doc = Document()

    # デフォルトフォント設定
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME_EN
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_JA)

    # 見出しスタイル
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = FONT_NAME_EN
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME_JA)
        hs.font.color.rgb = COLOR_DARK

    # ページ設定 (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # タイトル
    meta = data.get("meta", {})
    code = meta.get("company_code", "")
    filename = meta.get("filename", "")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    title_p.paragraph_format.space_after = Pt(12)
    run = title_p.add_run("経営分析レポート")
    fmt_run(run, size=24, bold=True, color=COLOR_DARK)

    if code:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_after = Pt(6)
        run = sub_p.add_run(f"企業コード: {code}")
        fmt_run(run, size=14, color=COLOR_ACCENT)

    if filename:
        sub_p2 = doc.add_paragraph()
        sub_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p2.paragraph_format.space_after = Pt(40)
        run = sub_p2.add_run(f"ソース: {filename}")
        fmt_run(run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    doc.add_page_break()

    # セクション描画
    for sec in data.get("sections", []):
        renderer = SECTION_RENDERERS.get(sec["id"])
        if renderer:
            renderer(doc, sec)
        else:
            print(f"[WARN] Unknown section id: {sec['id']}")

    # 締めの挨拶
    render_closing(doc)

    return doc


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="JSON → DOCX 変換")
    parser.add_argument("-i", "--input", type=str, required=True, help="入力 final_report_*.json のパス")
    parser.add_argument("-o", "--output", type=str, default=None, help="出力 .docx のパス")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    code = data.get("meta", {}).get("company_code", "unknown")
    output_path = args.output or os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
        "data", "final-output", "docx-per-company", f"report_{code}.docx",
    )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = build_document(data)
    doc.save(output_path)

    print(f"Input:  {args.input}")
    print(f"Output: {output_path}")
    print("Done.")


if __name__ == "__main__":
    main()
